#!/usr/bin/env python
"""Generate MS Word CRFs from CDASH metadata workbooks."""

import argparse
import pathlib

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _add_page_field(paragraph):
    """Insert a Word page number field into *paragraph*."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _add_bottom_border(cell) -> None:
    """Add a thin bottom border to *cell*."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "auto")


def _shade(cell, color: str) -> None:
    """Shade *cell* with *color* (hex RGB)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def _add_checkbox(paragraph) -> None:
    """Insert a checkbox content control into *paragraph*."""
    sdt = OxmlElement("w:sdt")
    pr = OxmlElement("w:sdtPr")
    cb = OxmlElement("w14:checkbox")
    pr.append(cb)
    content = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    content.append(r)
    sdt.append(pr)
    sdt.append(content)
    paragraph._p.append(sdt)


def _add_date_picker(paragraph) -> None:
    """Insert a date picker content control into *paragraph*."""
    sdt = OxmlElement("w:sdt")
    pr = OxmlElement("w:sdtPr")
    dt = OxmlElement("w14:date")
    pr.append(dt)
    content = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = ""
    r.append(t)
    content.append(r)
    sdt.append(pr)
    sdt.append(content)
    paragraph._p.append(sdt)


def _add_underline_entry(paragraph, length: int) -> None:
    run = paragraph.add_run(" " * length)
    run.font.underline = True


def load_ig(ig_path: str) -> pd.DataFrame:
    """Load and normalise the *Variables* worksheet."""
    ig_df = pd.read_excel(ig_path, sheet_name="Variables", engine="openpyxl")
    ig_df = ig_df[~ig_df["Domain"].isna()].copy()
    ig_df["Display Label"] = ig_df["Question Text"].fillna(
        ig_df["CDASHIG Variable Label"]
    )
    ig_df.rename(
        columns={
            "CDASHIG Variable": "Variable",
            "Variable Order": "Order",
            "Case Report Form Completion Instructions": "CRF Instructions",
            "CDISC CT Codelist Submission Values(s), Subset Submission Value(s)": "CT Values",
            "CDISC CT Codelist Code(s), Subset Codes(s)": "CT Codes",
        },
        inplace=True,
    )
    return ig_df


def build_domain_crf(
    domain_df: pd.DataFrame, domain: str, out_dir: pathlib.Path
) -> None:
    """Create a Word document for a single CDASH domain."""

    document = Document()

    # Use landscape orientation for readability
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Standardise font
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Header placeholders for protocol metadata
    hdr_width = section.page_width - section.left_margin - section.right_margin
    hdr_tbl = section.header.add_table(rows=1, cols=3, width=hdr_width)
    hdr_tbl.style = "Table Grid"
    hdr_cells = hdr_tbl.rows[0].cells
    hdr_cells[0].text = "Protocol ID: __________"
    hdr_cells[1].text = "Site Code: __________"
    hdr_cells[2].text = "Subject ID: __________"

    # Footer with version and automatic page numbering
    f_p = section.footer.add_paragraph("CRF Version 1.0 ")
    _add_page_field(f_p)
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_heading(f"{domain} Domain CRF", level=1)

    # Add a table with extra metadata columns to provide more context
    table = document.add_table(rows=1, cols=7, style="Table Grid")
    table.autofit = False
    total_width = section.page_width - section.left_margin - section.right_margin
    col_width = int(total_width / 7)
    for col in table.columns:
        col.width = col_width

    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "Label / Question"
    hdr[2].text = "Type"
    hdr[3].text = "Controlled Terminology"
    hdr[4].text = "Data Entry"
    hdr[5].text = "Instructions"
    hdr[6].text = "Required"
    _shade(hdr[6], "FFC7CE")

    ct_legend: dict[str, int] = {}
    footnotes: dict[str, int] = {}

    for idx, row in enumerate(domain_df.sort_values("Order").iterrows(), start=1):
        _, row = row
        cells = table.add_row().cells
        cells[0].text = row["Variable"]
        cells[1].text = str(row["Display Label"])
        cells[2].text = str(row.get("Type", ""))

        ct_val = row.get("CT Values")
        ct_code = row.get("CT Codes")
        if pd.notna(ct_val):
            ct = str(ct_val)
        elif pd.notna(ct_code):
            ct = str(ct_code)
        else:
            ct = ""

        if len(ct) > 40:
            idx_ct = ct_legend.setdefault(ct, len(ct_legend) + 1)
            cells[3].text = f"\u2020{idx_ct}"
        else:
            cells[3].text = ct

        entry_para = cells[4].paragraphs[0]
        label_lower = str(row.get("Display Label", "")).lower()
        var_upper = row["Variable"].upper()
        if (
            "date" in label_lower
            or var_upper.endswith("DT")
            or var_upper.endswith("DAT")
        ):
            _add_date_picker(entry_para)
        elif ct and len(str(ct).split(";")) <= 4:
            for i, val in enumerate(str(ct).split(";")):
                if i:
                    entry_para.add_run(" ")
                _add_checkbox(entry_para)
                entry_para.add_run(val.strip())
        else:
            expected = max(len(str(ct)), 10)
            _add_underline_entry(entry_para, expected)

        instructions = []
        if pd.notna(row.get("CRF Instructions")):
            instructions.append(str(row.get("CRF Instructions")))
        impl_note = row.get("Implementation Notes")
        if pd.notna(impl_note):
            impl_note = str(impl_note)
            if len(impl_note) > 60:
                fn_idx = footnotes.setdefault(impl_note, len(footnotes) + 1)
                instructions.append(f"[{fn_idx}]")
            else:
                instructions.append(impl_note)

            if any(t in impl_note.lower() for t in ["if ", "derive", "origin"]):
                instructions.append("Validate dependencies across domains")
        if (
            "date" in label_lower
            or var_upper.endswith("DT")
            or var_upper.endswith("DAT")
        ):
            instructions.append("Format: dd/mm/yyyy")

        instr_para = cells[5].paragraphs[0]
        for i_ins, item in enumerate(instructions):
            run = instr_para.add_run(item)
            run.italic = True
            if i_ins < len(instructions) - 1:
                instr_para.add_run("\n")

        req_cell = cells[6]
        req_text = row.get("CRF Instructions")
        if isinstance(req_text, str) and any(k in req_text.lower() for k in ["required", "mandatory"]):
            req_cell.text = "Yes"
        else:
            req_cell.text = ""


        if idx % 3 == 0:
            for c in cells:
                _add_bottom_border(c)

    if footnotes:
        document.add_heading("Footnotes", level=2)
        for text, num in footnotes.items():
            p = document.add_paragraph()
            p.add_run(f"[{num}] ").bold = True
            p.add_run(text)

    if ct_legend:
        document.add_page_break()
        document.add_heading("Controlled Terminology legend", level=2)
        legend = document.add_table(rows=len(ct_legend) + 1, cols=2, style="Table Grid")
        legend.cell(0, 0).text = "Symbol"
        legend.cell(0, 1).text = "Controlled Terminology"
        for ct_text, idx_ct in ct_legend.items():
            row_ct = legend.add_row().cells
            row_ct[0].text = f"\u2020{idx_ct}"
            row_ct[1].text = ct_text

    out_path = out_dir / f"{domain}_CRF.docx"
    document.save(out_path)
    print(f"\u2713 Saved {out_path.relative_to(out_dir.parent)}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate Word CRFs from CDASH metadata spreadsheets."
    )
    parser.add_argument(
        "--model",
        required=True,
        help="Path to CDASH_Model_v1.3.xlsx (reserved for future use)",
    )
    parser.add_argument("--ig", required=True, help="Path to CDASHIG_v2.3.xlsx")
    parser.add_argument(
        "--out",
        default="crfs",
        help="Directory to save the generated Word documents",
    )
    parser.add_argument(
        "--domains",
        nargs="*",
        metavar="DOMAIN",
        help="Optional list of domains to generate (e.g. AE CM)",
    )
    args = parser.parse_args()

    out_dir = pathlib.Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    ig_df = load_ig(args.ig)

    target_domains = args.domains or ig_df["Domain"].unique()
    for dom in target_domains:
        dom_df = ig_df[ig_df["Domain"] == dom]
        if dom_df.empty:
            print(f"\u26a0 Domain {dom} not found in IG \u2013 skipped")
            continue
        build_domain_crf(dom_df, dom, out_dir)


if __name__ == "__main__":
    main()
