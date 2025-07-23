from pathlib import Path
from typing import Sequence

import docx

from crfgen.schema import Form

from .registry import register


@register("docx")
def export_docx(forms: Sequence[Form], outdir: Path) -> None:
    doc = docx.Document()
    for form in forms:
        doc.add_heading(form.title, level=1)
        for fld in form.fields:
            doc.add_paragraph(f"{fld.prompt} ({fld.oid})")
    doc.save(outdir / "forms.docx")
