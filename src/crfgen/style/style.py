from docx import Document
from docx.shared import Pt, RGBColor

def apply_styles(doc: Document, styles: dict):
    """
    Applies a set of styles to a docx document.
    """
    if not styles:
        return

    # Apply font styles
    font_styles = styles.get("font")
    if font_styles:
        for style in doc.styles:
            if hasattr(style, "font"):
                font = style.font
                font.name = font_styles.get("name", font.name)
                if font.size:
                    font.size = Pt(font_styles.get("size", font.size.pt))
                else:
                    font.size = Pt(font_styles.get("size", 10))

    # Add a logo to the header
    logo_path = styles.get("logo_path")
    if logo_path:
        header = doc.sections[0].header
        header.paragraphs[0].add_run().add_picture(logo_path)

    # Apply color styles
    colors = styles.get("colors")
    if colors:
        # Example of changing the color of table headers
        # This will need to be more sophisticated and tied to the
        # actual table generation logic.
        pass
