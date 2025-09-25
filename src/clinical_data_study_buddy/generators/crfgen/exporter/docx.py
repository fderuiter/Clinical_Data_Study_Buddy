"""
This module provides the functionality to export CRF (Case Report Form) data
to a Microsoft Word (.docx) file.
"""

from pathlib import Path
from typing import Sequence

import docx

from clinical_data_study_buddy.core.models.schema import Form
from clinical_data_study_buddy.generators.crfgen.style.style import apply_styles

from .registry import register


@register("docx")
def export_docx(
    forms: Sequence[Form], outdir: Path, style: dict = None, output_filename: str = None
) -> None:
    """
    Exports a sequence of Form objects to a .docx file.

    This function creates a Word document, applies styles if provided, and then
    populates the document with data from the forms.

    Args:
        forms (Sequence[Form]): A sequence of Form objects to be exported.
        outdir (Path): The output directory where the .docx file will be saved.
        style (dict, optional): A dictionary defining the styles to be applied
                                to the document. Defaults to None.
        output_filename (str, optional): The desired name for the output file.
                                         If not provided, a name will be generated
                                         based on the first form's ID.
    """
    doc = docx.Document()
    if style:
        apply_styles(doc, style)

    for form in forms:
        doc.add_heading(form.title, level=1)
        if hasattr(form, "data") and form.data:
            if isinstance(form.data, list) and form.data:
                table = doc.add_table(rows=1, cols=len(form.data[0]))
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(form.data[0].keys()):
                    hdr_cells[i].text = key
                for item in form.data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(item.values()):
                        row_cells[i].text = str(value)
            else:
                doc.add_paragraph("No data available for this table.")
        else:
            if hasattr(form, "fields") and form.fields:
                for fld in form.fields:
                    doc.add_paragraph(f"{fld.prompt} ({fld.oid})")

    if output_filename:
        doc.save(outdir / output_filename)
    else:
        doc.save(outdir / f"{forms[0].id}.docx")
