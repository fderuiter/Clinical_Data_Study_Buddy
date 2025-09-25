"""
This module provides the core functionality for building Case Report Forms (CRFs)
in Microsoft Word format based on CDASH IG standards.

It includes functions for:
- Fetching CDASHIG metadata from the CDISC Library API.
- Creating and styling various components of a Word document, such as headers,
  footers, tables, and form controls.
- Assembling the complete CRF for a given domain.
"""

import os
import pathlib
from typing import Any, Dict, List, Tuple

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

from cdisc_library_client.api.cdash_implementation_guide_cdashig import (
    get_mdr_cdashig_version_domains,
    get_mdr_cdashig_version_domains_domain_fields,
)
from cdisc_library_client.client import AuthenticatedClient

###############################################################################
# Domain‑to‑category mapping
###############################################################################

DOMAIN_INFO: Dict[str, Tuple[str, str]] = {
    # Interventions
    "AG": ("Interventions", "Procedure Agents"),
    "CM": ("Interventions", "Concomitant / Prior Medications"),
    "EC": ("Interventions", "Exposure as Collected"),
    "EX": ("Interventions", "Exposure"),
    "ML": ("Interventions", "Meal Data"),
    "PR": ("Interventions", "Procedures"),
    "SU": ("Interventions", "Substance Use"),
    # Events
    "AE": ("Events", "Adverse Events"),
    "CE": ("Events", "Clinical Events"),
    "DS": ("Events", "Disposition"),
    "DV": ("Events", "Protocol Deviations"),
    "HO": ("Events", "Healthcare Encounters"),
    "MH": ("Events", "Medical History"),
    "SA": ("Events", "Serious Adverse Events"),
    # Findings
    "CP": ("Findings", "Cell Phenotype Findings"),
    "CV": ("Findings", "Cardiovascular System Findings"),
    "DA": ("Findings", "Product Accountability"),
    "DD": ("Findings", "Death Details"),
    "ED": ("Findings", "Central Reading"),
    "GF": ("Findings", "Genomics Findings"),
    "IE": ("Findings", "Inclusion / Exclusion Criteria Not Met"),
    "LB": ("Findings", "Laboratory Test Results"),
    "MB": ("Findings", "Microbiology Specimen"),
    "MI": ("Findings", "Microscopic Findings"),
    "MK": ("Findings", "Musculoskeletal System Findings"),
    "MS": ("Findings", "Microbiology Susceptibility"),
    "NV": ("Findings", "Nervous System Findings"),
    "OE": ("Findings", "Ophthalmic Examinations"),
    "PC": ("Findings", "Pharmacokinetics Concentrations"),
    "PE": ("Findings", "Physical Examination"),
    "RE": ("Findings", "Respiratory System Findings"),
    "RP": ("Findings", "Reproductive System Findings"),
    "RS": ("Findings", "Disease Response & Clinical Classification"),
    "SC": ("Findings", "Subject Characteristics"),
    "TR": ("Findings", "Tumor / Lesion Results"),
    "TU": ("Findings", "Tumor / Lesion Identification"),
    "UR": ("Findings", "Urinary System Findings"),
    "VS": ("Findings", "Vital Signs"),
    # Findings‑About
    "FA": ("Findings About", "Findings About Events or Interventions"),
    "SR": ("Findings About", "Skin Response"),
    # Special Purpose
    "CO": ("Special Purpose", "Comments"),
    "DM": ("Special Purpose", "Demographics"),
}


REPEATING_DOMAINS = {"AE", "CE", "CM", "DS", "DV", "HO", "MH", "SA"}


def get_domain_info(domain: str) -> Tuple[str, str]:
    """
    Returns the category and full title for a given domain code.

    Args:
        domain (str): The two-letter domain code (e.g., "AE").

    Returns:
        Tuple[str, str]: A tuple containing the category and the full title
                         of the domain. If the domain is not found, it returns
                         ("Unknown", domain).
    """
    try:
        return DOMAIN_INFO[domain.upper()]
    except KeyError:
        return ("Unknown", domain)


###############################################################################
# Low‑level helpers
###############################################################################


def _add_page_field(paragraph):
    """
    Inserts a Word PAGE field into a paragraph.

    This function modifies the paragraph in-place to include a field that
    will display the current page number.

    Args:
        paragraph: The docx.paragraph.Paragraph object to which the page
                   field will be added.
    """
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _set_cell_shading(cell, color_hex: str = "4F81BD"):
    """
    Shades a table cell's background with a specified color.

    Args:
        cell: The docx.table._Cell object to be shaded.
        color_hex (str): The RGB hex string for the color (without the '#').
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove existing shading if any
    for shd in tc_pr.findall("w:shd", tc_pr.nsmap):
        tc_pr.remove(shd)
    # Add new shading element
    shd_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tc_pr.append(shd_elm)


def _add_bottom_border(cell) -> None:
    """
    Adds a thin bottom border to a table cell.

    Args:
        cell: The docx.table._Cell object to which the border will be added.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "auto")


def _add_checkbox(paragraph) -> None:
    """
    Inserts a checkbox content control into a paragraph.

    Args:
        paragraph: The docx.paragraph.Paragraph object where the checkbox
                   will be inserted.
    """
    sdt = OxmlElement("w:sdt")
    pr = OxmlElement("w:sdtPr")
    cb = OxmlElement("w14:checkbox")
    pr.append(cb)
    content = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    content.append(r)
    sdt.append(pr)
    sdt.append(content)
    paragraph._p.append(sdt)


def _add_date_picker(paragraph) -> None:
    """
    Inserts a date picker content control into a paragraph.

    Args:
        paragraph: The docx.paragraph.Paragraph object where the date picker
                   will be inserted.
    """
    sdt = OxmlElement("w:sdt")
    pr = OxmlElement("w:sdtPr")
    dt = OxmlElement("w14:date")
    pr.append(dt)
    content = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = ""
    r.append(t)
    content.append(r)
    sdt.append(pr)
    sdt.append(content)
    paragraph._p.append(sdt)


def _add_underline_entry(paragraph, length: int) -> None:
    """
    Adds an underlined space for manual data entry.

    Args:
        paragraph: The docx.paragraph.Paragraph object where the entry
                   line will be added.
        length (int): The number of space characters to underline.
    """
    run = paragraph.add_run(" " * length)
    run.font.underline = True


def _style_header_cell(cell):
    """
    Applies bold, white font styling to a header cell.

    Args:
        cell: The docx.table._Cell object to be styled.
    """
    para = cell.paragraphs[0]
    run = para.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


###############################################################################
# Data I/O helpers
###############################################################################


def get_cdashig_variables_from_api(ig_version: str) -> pd.DataFrame:
    """
    Loads and normalizes CDASHIG variables from the CDISC Library API.

    Args:
        ig_version (str): The version of the CDASHIG to fetch (e.g., "v2.3").

    Returns:
        pd.DataFrame: A DataFrame containing the normalized CDASHIG variables.

    Raises:
        ValueError: If the CDISC_PRIMARY_KEY environment variable is not set.
    """
    api_key = os.environ.get("CDISC_PRIMARY_KEY")
    if not api_key:
        raise ValueError("CDISC_PRIMARY_KEY environment variable not set.")

    client = AuthenticatedClient(
        base_url="https://library.cdisc.org/api", token=api_key
    )

    all_variables = []

    # Get all domains for the given CDASHIG version
    domains_response = get_mdr_cdashig_version_domains.sync(
        client=client, version=ig_version
    )
    if (
        not domains_response
        or not domains_response.field_links
        or not domains_response.field_links.domains
    ):
        print(f"Warning: No domains found for CDASHIG version {ig_version}")
        return pd.DataFrame()

    for domain_item in domains_response.field_links.domains:
        domain_name = domain_item.href.split("/")[-1]
        if not domain_name:
            continue

        # Get all fields for the given domain
        page = 1
        while True:
            fields_response = get_mdr_cdashig_version_domains_domain_fields.sync(
                client=client,
                version=ig_version,
                domain=domain_name,
                page=page,
                page_size=100,
            )
            if (
                not fields_response
                or not fields_response.field_links
                or not fields_response.field_links.fields
            ):
                break

            from cdisc_library_client.api.cdash_implementation_guide_cdashig import (
                get_mdr_cdashig_version_domains_domain_fields_field,
            )

            for field_ref in fields_response.field_links.fields:
                field_name = field_ref.href.split("/")[-1]
                field_details = (
                    get_mdr_cdashig_version_domains_domain_fields_field.sync(
                        client=client,
                        version=ig_version,
                        domain=domain_name,
                        field=field_name,
                    )
                )
                if not field_details:
                    continue

                variable_data = {
                    "Domain": domain_name,
                    "Variable": field_details.name,
                    "Order": field_details.ordinal,
                    "Display Label": field_details.prompt or field_details.label,
                    "CRF Instructions": field_details.completion_instructions,
                    "Type": field_details.simple_datatype,
                    "CT Values": (
                        "; ".join(
                            field_details.additional_properties.get(
                                "codelistSubmissionValues", []
                            )
                        )
                        if field_details.additional_properties
                        and field_details.additional_properties.get(
                            "codelistSubmissionValues"
                        )
                        else None
                    ),
                    "CT Codes": None,  # Not available in this endpoint
                    "Implementation Notes": field_details.implementation_notes,
                }
                all_variables.append(variable_data)

            page += 1

    df = pd.DataFrame(all_variables)
    return df


def load_ig(ig_version: str) -> pd.DataFrame:
    """
    Loads and normalizes CDASHIG variables from the CDISC Library API.

    This function serves as a wrapper around get_cdashig_variables_from_api.

    Args:
        ig_version (str): The version of the CDASHIG to fetch.

    Returns:
        pd.DataFrame: A DataFrame containing the CDASHIG variables.
    """
    return get_cdashig_variables_from_api(ig_version)


###############################################################################
# Core CRF builder
###############################################################################


def _create_header(section, config, full_title):
    """
    Creates the page header for the CRF document.

    Args:
        section: The docx.section.Section object for the document.
        config (dict): A dictionary containing study metadata.
        full_title (str): The full title of the CRF.
    """
    header = section.header
    hdr_tbl = header.add_table(rows=2, cols=2, width=section.page_width)
    hdr_tbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_tbl.autofit = False
    hdr_tbl.repeat_rows = 0

    sponsor_cell, title_cell = hdr_tbl.rows[0].cells
    meta_cell_L, meta_cell_R = hdr_tbl.rows[1].cells

    # Row-0: sponsor block & CRF title block
    sponsor_cell.text = config.get("study_metadata", {}).get(
        "sponsor_name", "Sponsor Study Name"
    )
    title_cell.text = full_title
    sponsor_cell.width = title_cell.width = section.page_width / 2

    header_color = config.get("styling", {}).get("header_color", "1F1F1F")
    for c in (sponsor_cell, title_cell):
        _set_cell_shading(c, header_color)
        _style_header_cell(c)

    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_para.runs[0].font.size = Pt(14)

    # Row-1: Subject meta-data placeholders
    protocol_id = config.get("study_metadata", {}).get("protocol_id", "_____‑___‑___")
    meta_cell_L.text = f"Subject ID: {protocol_id}    SITE #: ___"
    meta_cell_R.text = "Initials: ___ ___ ___"
    for c in (meta_cell_L, meta_cell_R):
        _set_cell_shading(c, "3F3F3F")
        _style_header_cell(c)


def _create_footer(section, config, full_title):
    """
    Creates the page footer for the CRF document.

    Args:
        section: The docx.section.Section object for the document.
        config (dict): A dictionary containing the version label.
        full_title (str): The full title of the CRF.
    """
    footer = section.footer
    version_label = config.get("version_label", "Version 1.0 DRAFT")
    f_left = footer.add_paragraph(f"{full_title}, {version_label}")
    f_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    f_left.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    f_right = footer.add_paragraph()
    f_right.paragraph_format.right_indent = Pt(0)
    _add_page_field(f_right)
    f_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_right.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _create_admin_section(document, full_title, config):
    """
    Creates Section A (Administrative) of the CRF.

    Args:
        document: The docx.Document object.
        full_title (str): The full title of the CRF.
        config (dict): A dictionary containing configuration settings.
    """
    document.add_paragraph()
    secA_tbl = document.add_table(rows=3, cols=2, style="Table Grid")
    secA_tbl.autofit = False
    secA_tbl.allow_autofit = False

    hdr_row = secA_tbl.rows[0]
    hdr_cell = hdr_row.cells[0]
    hdr_cell.merge(hdr_row.cells[1])
    hdr_cell.text = "SECTION A  ADMINISTRATIVE"
    section_header_color = config.get("styling", {}).get(
        "section_header_color", "8064A2"
    )
    _set_cell_shading(hdr_cell, section_header_color)
    _style_header_cell(hdr_cell)

    secA_tbl.rows[1].cells[0].text = f"Was {full_title.lower()} completed?"
    p = secA_tbl.rows[1].cells[1].paragraphs[0]
    _add_checkbox(p)
    p.add_run(" No (Complete protocol deviation form)    ")
    _add_checkbox(p)
    p.add_run(" Yes")

    secA_tbl.rows[2].cells[0].text = "Date of assessment:"
    secA_tbl.rows[2].cells[1].text = "__|__|____|____|    DD-MMM-YYYY"


def _create_variables_table(
    document,
    section,
    domain_df,
    domain: str,
    config: dict,
    fda_adverse_events: List[Dict[str, Any]] = None,
):
    """
    Creates Section B (Domain Variables) of the CRF.

    This function builds the main table of variables for the domain, including
    data entry controls and footnotes.

    Args:
        document: The docx.Document object.
        section: The docx.section.Section object for the document.
        domain_df (pd.DataFrame): A DataFrame containing the variables for the domain.
        domain (str): The two-letter domain code.
        config (dict): A dictionary containing configuration settings.
        fda_adverse_events (List[Dict[str, Any]], optional): A list of adverse
            event terms from OpenFDA. Defaults to None.
    """
    document.add_paragraph()
    var_tbl = document.add_table(rows=1, cols=6, style="Table Grid")
    var_tbl.autofit = False
    total_width = section.page_width - section.left_margin - section.right_margin
    col_width = int(total_width / 6)
    for col in var_tbl.columns:
        col.width = col_width

    hdr_cells = var_tbl.rows[0].cells
    col_titles = [
        "Variable",
        "Label / Question",
        "Type",
        "Controlled Terminology",
        "Data Entry",
        "Instructions",
    ]
    table_header_color = config.get("styling", {}).get("table_header_color", "4F81BD")
    for idx, title in enumerate(col_titles):
        hdr_cells[idx].text = title
        _set_cell_shading(hdr_cells[idx], table_header_color)
        _style_header_cell(hdr_cells[idx])

    ct_legend: dict[str, int] = {}
    footnotes: dict[str, int] = {}

    # Data rows ordered by the "Variable Order" column
    for idx, (_, row) in enumerate(domain_df.sort_values("Order").iterrows(), start=1):
        cells = var_tbl.add_row().cells
        # 0 Variable name
        cells[0].text = row["Variable"]

        # 1 Label / Question
        cells[1].text = str(row["Display Label"])

        # 2 Data type (where available)
        cells[2].text = str(row.get("Type", ""))

        # 3 Controlled terminology – prefer values over codes
        ct_val = row.get("CT Values")
        ct_code = row.get("CT Codes")
        ct = (
            str(ct_val)
            if pd.notna(ct_val)
            else str(ct_code) if pd.notna(ct_code) else ""
        )
        if len(ct) > 40:
            idx_ct = ct_legend.setdefault(ct, len(ct_legend) + 1)
            cells[3].text = f"\u2020{idx_ct}"
        else:
            cells[3].text = ct

        # 4 Data entry placeholder (smart entry controls)
        entry_para = cells[4].paragraphs[0]
        label_lower = str(row.get("Display Label", "")).lower()
        var_upper = row["Variable"].upper()
        if "date" in label_lower or var_upper.endswith(("DT", "DAT")):
            _add_date_picker(entry_para)
        elif ct and len(str(ct).split(";")) <= 4:
            for i, val in enumerate(str(ct).split(";")):
                if i:
                    entry_para.add_run(" ")
                _add_checkbox(entry_para)
                entry_para.add_run(val.strip())
        else:
            expected = max(len(str(ct)), 10)
            _add_underline_entry(entry_para, expected)

        # 5 Instructions (italic, stacked if multiple)
        instructions = []
        if pd.notna(row.get("CRF Instructions")):
            instructions.append(str(row.get("CRF Instructions")))
        impl_note = row.get("Implementation Notes")
        if pd.notna(impl_note):
            impl_note = str(impl_note)
            if len(impl_note) > 60:
                fn_idx = footnotes.setdefault(impl_note, len(footnotes) + 1)
                instructions.append(f"[{fn_idx}]")
            else:
                instructions.append(impl_note)

            if any(t in impl_note.lower() for t in ["if ", "derive", "origin"]):
                instructions.append("Validate dependencies across domains")

        # Auto-detect date fields and add a formatting hint
        if "date" in label_lower or var_upper.endswith(("DT", "DAT")):
            instructions.append("Format: dd/mm/yyyy")

        instr_para = cells[5].paragraphs[0]
        for i_ins, item in enumerate(instructions):
            run = instr_para.add_run(item)
            run.italic = True
            if i_ins < len(instructions) - 1:
                instr_para.add_run("\n")

        if idx % 3 == 0:
            for c in cells:
                _add_bottom_border(c)

    if domain.upper() in REPEATING_DOMAINS:
        instruction_row = var_tbl.add_row()
        instruction_cell = instruction_row.cells[0]
        instruction_cell.merge(instruction_row.cells[-1])
        p = instruction_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Right-click -> Insert -> Insert Rows Below to add more entries."
        )
        run.italic = True
        run.font.size = Pt(9)

    if domain.upper() == "AE" and fda_adverse_events:
        document.add_page_break()
        document.add_heading("Suggested Adverse Event Terms from OpenFDA", level=2)
        fda_tbl = document.add_table(rows=1, cols=1, style="Table Grid")
        hdr_cell = fda_tbl.rows[0].cells[0]
        hdr_cell.text = "Reaction Term"
        _set_cell_shading(hdr_cell, table_header_color)
        _style_header_cell(hdr_cell)
        for event in fda_adverse_events:
            row_cells = fda_tbl.add_row().cells
            row_cells[0].text = event["reaction_term"]

    if footnotes:
        document.add_heading("Footnotes", level=2)
        for text, num in footnotes.items():
            p = document.add_paragraph()
            p.add_run(f"[{num}] ").bold = True
            p.add_run(text)

    if ct_legend:
        document.add_page_break()
        document.add_heading("Controlled Terminology legend", level=2)
        legend = document.add_table(rows=len(ct_legend) + 1, cols=2, style="Table Grid")
        legend.cell(0, 0).text = "Symbol"
        legend.cell(0, 1).text = "Controlled Terminology"
        for ct_text, idx_ct in ct_legend.items():
            row_ct = legend.add_row().cells
            row_ct[0].text = f"\u2020{idx_ct}"
            row_ct[1].text = ct_text


def build_domain_crf(
    domain_df: pd.DataFrame,
    domain: str,
    out_dir: pathlib.Path,
    config: dict,
    fda_adverse_events: List[Dict[str, Any]] = None,
) -> None:
    """
    Builds a Word document for a single CDASH domain and saves it to disk.

    Args:
        domain_df (pd.DataFrame): A DataFrame containing the variables for the domain.
        domain (str): The two-letter domain code.
        out_dir (pathlib.Path): The directory where the generated Word document
                                will be saved.
        config (dict): A dictionary containing configuration settings, including
                       study metadata and version information.
        fda_adverse_events (List[Dict[str, Any]], optional): A list of adverse
            event terms from OpenFDA. Defaults to None.
    """

    category, full_title = get_domain_info(domain)

    # ---------------------------------------------------------------------
    #  Document meta & base formatting
    # ---------------------------------------------------------------------
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Uniform font for entire document
    style = document.styles["Normal"]
    styling = config.get("styling", {})
    style.font.name = styling.get("font_name", "Arial")
    style.font.size = Pt(styling.get("font_size", 10))

    # ---------------------------------------------------------------------
    #  Create document components
    # ---------------------------------------------------------------------
    _create_header(section, config, full_title)
    _create_footer(section, config, full_title)
    _create_admin_section(document, full_title, config)
    _create_variables_table(
        document,
        section,
        domain_df,
        domain,
        config,
        fda_adverse_events=fda_adverse_events,
    )

    # ---------------------------------------------------------------------
    #  Save document
    # ---------------------------------------------------------------------
    safe_title = full_title.replace(" / ", "_").replace(" ", "_")
    out_path = out_dir / f"{domain}_{safe_title}_CRF.docx"
    document.save(out_path)
    print(f"\u2713 Saved {out_path.relative_to(out_dir.parent)}")
