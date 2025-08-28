import json
from docx import Document

def generate_sdrg(crf_json_path, study_config_path, output_path):
    """
    Generates a Study Data Reviewer's Guide (SDRG) document in .docx format.

    This function creates a Word document that serves as a guide for reviewers of study data.
    It includes sections for introduction, protocol description, list of included documents,
    data collection and processing, SDTM datasets, and more. The content is populated based
    on the provided CRF JSON and study configuration.

    :param crf_json_path: Path to the canonical CRF JSON file.
    :param study_config_path: Path to the study-specific configuration file.
    :param output_path: Path to save the generated Word document.
    """
    # Load data from crf.json and study_config.json
    with open(crf_json_path, 'r') as f:
        crf_data = json.load(f)
    with open(study_config_path, 'r') as f:
        study_config = json.load(f)

    # Create a new Word document
    document = Document()

    # Add title
    document.add_heading('Study Data Reviewer\'s Guide', level=0)

    # Section 1: Introduction
    document.add_heading('1. Introduction', level=1)
    document.add_paragraph(f"This document describes the study data for {study_config.get('study_id', 'this study')}.")
    document.add_paragraph("This guide is intended to assist the reviewer in understanding the structure and content of the study data.")

    # Section 2: Protocol Description
    document.add_heading('2. Protocol Description', level=1)
    document.add_paragraph(f"Protocol: {study_config.get('protocol_id', 'N/A')}")
    document.add_paragraph(f"Protocol Title: {study_config.get('protocol_title', 'N/A')}")

    # Section 3: List of Included Documents
    document.add_heading('3. List of Included Documents', level=1)
    document.add_paragraph("This submission includes the following documents:")
    document.add_paragraph("- SDRG (this document)")
    document.add_paragraph("- define.xml")
    document.add_paragraph("- Annotated CRF")

    # Section 4: Data Collection and Processing
    document.add_heading('4. Data Collection and Processing', level=1)
    document.add_paragraph("This section describes how the data was collected and processed.")

    # Section 5: SDTM Datasets
    document.add_heading('5. SDTM Datasets', level=1)
    document.add_paragraph("The following SDTM datasets are included in this submission:")
    # This section would be populated with information about the SDTM datasets
    document.add_paragraph("- DM (Demographics)")
    document.add_paragraph("- AE (Adverse Events)")

    # Section 6: Data Conformance
    document.add_heading('6. Data Conformance', level=1)
    document.add_paragraph("The study data were checked for conformance with the CDISC SDTM standard.")
    document.add_paragraph("Any conformance issues are documented in the data definition file (define.xml).")

    # Section 7: Appendices
    document.add_heading('7. Appendices', level=1)
    document.add_paragraph("This section contains any appendices.")

    # Save the document
    document.save(output_path)
    print(f"SDRG document generated at: {output_path}")
