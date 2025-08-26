import json
from docx import Document

def generate_adrg(crf_json_path, study_config_path, output_path):
    """
    Generates an Analysis Data Reviewer's Guide (ADRG) document.
    """
    # Load data from crf.json and study_config.json
    with open(crf_json_path, 'r') as f:
        crf_data = json.load(f)
    with open(study_config_path, 'r') as f:
        study_config = json.load(f)

    # Create a new Word document
    document = Document()

    # Add title
    document.add_heading('Analysis Data Reviewer\'s Guide', level=0)

    # Section 1: Introduction
    document.add_heading('1. Introduction', level=1)
    document.add_paragraph(f"This document describes the analysis datasets for {study_config.get('study_id', 'this study')}.")
    document.add_paragraph("This guide is intended to assist the reviewer in understanding the structure and content of the analysis datasets.")

    # Section 2: Protocol Description
    document.add_heading('2. Protocol Description', level=1)
    document.add_paragraph(f"Protocol: {study_config.get('protocol_id', 'N/A')}")
    document.add_paragraph(f"Protocol Title: {study_config.get('protocol_title', 'N/A')}")

    # Section 3: Analysis Datasets
    document.add_heading('3. Analysis Datasets', level=1)
    document.add_paragraph("The following analysis datasets are included in this submission:")
    # This section would be populated with information about the ADaM datasets
    # For now, we'll just add a placeholder
    document.add_paragraph("- ADSL (Subject-Level Analysis Dataset)")
    document.add_paragraph("- ADAE (Adverse Event Analysis Dataset)")

    # Section 4: Data Conformance
    document.add_heading('4. Data Conformance', level=1)
    document.add_paragraph("The analysis datasets were checked for conformance with the CDISC ADaM standard.")
    document.add_paragraph("Any conformance issues are documented in the data definition file (define.xml).")

    # Section 5: Program and Macro Catalog
    document.add_heading('5. Program and Macro Catalog', level=1)
    document.add_paragraph("This section provides a catalog of the programs and macros used to create the analysis datasets.")
    # This section would be populated with a list of programs and macros
    document.add_paragraph("- adsl.sas")
    document.add_paragraph("- adae.sas")

    # Section 6: Data Listings
    document.add_heading('6. Data Listings', level=1)
    document.add_paragraph("This section provides information about the data listings included in this submission.")

    # Section 7: Figures and Tables
    document.add_heading('7. Figures and Tables', level=1)
    document.add_paragraph("This section provides information about the figures and tables included in this submission.")

    # Save the document
    document.save(output_path)
    print(f"ADRG document generated at: {output_path}")
