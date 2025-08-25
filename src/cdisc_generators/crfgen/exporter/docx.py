from pathlib import Path
from typing import Sequence

import docx

from cdisc_generators.crfgen.schema import Form
from cdisc_generators.crfgen.style.style import apply_styles
from .registry import register


@register("docx")
def export_docx(forms: Sequence[Form], outdir: Path, style: dict = None, output_filename: str = None) -> None:
    doc = docx.Document()
    if style:
        apply_styles(doc, style)

    for form in forms:
        doc.add_heading(form.title, level=1)
        if form.data:
            if isinstance(form.data, list) and form.data:
                table = doc.add_table(rows=1, cols=len(form.data[0]))
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(form.data[0].keys()):
                    hdr_cells[i].text = key
                for item in form.data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(item.values()):
                        row_cells[i].text = str(value)
            else:
                doc.add_paragraph("No data available for this table.")
        else:
            if hasattr(form, 'fields') and form.fields:
                for fld in form.fields:
                    doc.add_paragraph(f"{fld.prompt} ({fld.oid})")

    if output_filename:
        doc.save(outdir / output_filename)
    else:
        doc.save(outdir / f"{forms[0].id}.docx")
