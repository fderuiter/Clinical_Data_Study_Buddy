import pandas as pd
from unittest.mock import patch
from typer.testing import CliRunner
from clinical_data_study_buddy.cli.main import app
from docx import Document
from zipfile import ZipFile

runner = CliRunner()

def test_help():
    result = runner.invoke(app, ["generate", "cdash-crf", "--help"])
    assert result.exit_code == 0
    assert "Generates Word CRF (Case Report Form) shells" in result.stdout


@patch("clinical_data_study_buddy.core.generation_service.load_ig")
def test_generate(mock_load_ig, tmp_path):
    out_dir = tmp_path / "out"
    mock_df = pd.DataFrame([
        {
            "Domain": "AE", "Variable": "AETERM", "Order": 1,
            "Display Label": "Adverse Event Term",
            "CRF Instructions": "Report the term for the adverse event.",
            "Type": "Char", "CT Values": None, "CT Codes": None,
            "Implementation Notes": "If the event is serious, this is a very long implementation note that is definitely longer than sixty characters to ensure that the footnotes section is created."
        },
        {
            "Domain": "AE", "Variable": "AESEV", "Order": 2,
            "Display Label": "Severity",
            "CRF Instructions": "Report the severity of the adverse event.",
            "Type": "Char", "CT Values": "MILD; MODERATE; SEVERE", "CT Codes": None,
            "Implementation Notes": None
        },
        {
            "Domain": "AE", "Variable": "AESTDTC", "Order": 3,
            "Display Label": "Start Date",
            "CRF Instructions": "Report the start date of the adverse event.",
            "Type": "Datetime", "CT Values": None, "CT Codes": None,
            "Implementation Notes": None
        }
    ])
    mock_load_ig.return_value = mock_df

    result = runner.invoke(app, ["generate", "cdash-crf", "--ig-version", "v2.3", "--out", str(out_dir), "--domains", "AE"])
    assert result.exit_code == 0, result.stdout

    doc_path = out_dir / "AE_Adverse_Events_CRF.docx"
    assert doc_path.exists()

    doc = Document(doc_path)
    table = doc.tables[1]
    assert len(table.columns) == 6

    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Footnotes" in texts
    assert "[1]" in texts

    admin = doc.tables[0].cell(0, 0).text.replace("\xa0", " ")
    assert admin == "SECTION A  ADMINISTRATIVE"

    with ZipFile(doc_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        assert "w14:checkbox" in xml or "w14:date" in xml
        assert "Validate dependencies" in xml
        assert xml.count("<w:bottom") > 0

    assert len(doc.tables) == 2
    table = doc.tables[1]
    assert len(table.columns) == 6


@patch("clinical_data_study_buddy.core.generation_service.populate_ae_from_fda")
@patch("clinical_data_study_buddy.core.generation_service.load_ig")
def test_generate_with_openfda(mock_load_ig, mock_populate_ae, tmp_path):
    out_dir = tmp_path / "out"
    mock_df = pd.DataFrame([
        {"Domain": "AE", "Variable": "AETERM", "Order": 1, "Display Label": "Adverse Event Term", "CRF Instructions": "", "Type": "Char", "CT Values": None, "CT Codes": None, "Implementation Notes": None},
    ])
    mock_load_ig.return_value = mock_df
    mock_populate_ae.return_value = [{"reaction_term": "Headache"}, {"reaction_term": "Nausea"}]

    result = runner.invoke(app, ["generate", "cdash-crf", "--ig-version", "v2.3", "--out", str(out_dir), "--domains", "AE", "--openfda-drug-name", "TestDrug"])
    assert result.exit_code == 0, result.stdout

    doc_path = out_dir / "AE_Adverse_Events_CRF.docx"
    assert doc_path.exists()

    doc = Document(doc_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert "Suggested Adverse Event Terms from OpenFDA" in headings

    assert len(doc.tables) == 3
    fda_table = doc.tables[2]
    assert fda_table.cell(0, 0).text == "Reaction Term"
    assert fda_table.cell(1, 0).text == "Headache"
    assert fda_table.cell(2, 0).text == "Nausea"


@patch("clinical_data_study_buddy.core.generation_service.load_ig")
def test_generate_with_custom_styling(mock_load_ig, tmp_path):
    out_dir = tmp_path / "out"
    config_path = tmp_path / "config.yaml"
    with open(config_path, "w") as f:
        f.write("""styling:
  header_color: "FF0000"
  section_header_color: "00FF00"
  table_header_color: "0000FF"
  font_name: "Comic Sans MS"
  font_size: 12
""")
    mock_df = pd.DataFrame([
        {
            "Domain": "AE", "Variable": "AETERM", "Order": 1,
            "Display Label": "Adverse Event Term",
            "CRF Instructions": "Report the term for the adverse event.",
            "Type": "Char", "CT Values": None, "CT Codes": None,
            "Implementation Notes": "Test"
        }
    ])
    mock_load_ig.return_value = mock_df

    result = runner.invoke(app, ["generate", "cdash-crf", "--ig-version", "v2.3", "--out", str(out_dir), "--domains", "AE", "--config", str(config_path)])
    assert result.exit_code == 0, result.stdout

    doc_path = out_dir / "AE_Adverse_Events_CRF.docx"
    assert doc_path.exists()

    with ZipFile(doc_path) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
        assert 'w:fill="00FF00"' in doc_xml
        assert 'w:fill="0000FF"' in doc_xml
        header_xml = ""
        for name in zf.namelist():
            if "header" in name:
                header_xml = zf.read(name).decode("utf-8")
                break
        assert 'w:fill="FF0000"' in header_xml
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
        assert '<w:rFonts w:ascii="Comic Sans MS"' in styles_xml
        assert '<w:sz w:val="24"/>' in styles_xml
