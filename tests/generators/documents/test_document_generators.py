import pytest
import pathlib
from docx import Document
from clinical_data_study_buddy.generators.documents.adrg_generator import ADRGGenerator
from clinical_data_study_buddy.generators.documents.sdrg_generator import SDRGGenerator

@pytest.fixture
def study_config_path():
    return pathlib.Path("tests/fixtures/study_config.json")

def test_adrg_generator(study_config_path, tmp_path):
    output_path = tmp_path / "adrg.docx"
    generator = ADRGGenerator(study_config_path)
    generator.generate(output_path)

    assert output_path.exists()
    doc = Document(output_path)
    assert doc.paragraphs[0].text == "Analysis Data Reviewer's Guide"

def test_sdrg_generator(study_config_path, tmp_path):
    output_path = tmp_path / "sdrg.docx"
    generator = SDRGGenerator(study_config_path)
    generator.generate(output_path)

    assert output_path.exists()
    doc = Document(output_path)
    assert doc.paragraphs[0].text == "Study Data Reviewer's Guide"

def test_sdrg_generator_content(study_config_path, tmp_path):
    output_path = tmp_path / "sdrg.docx"
    generator = SDRGGenerator(study_config_path)
    generator.generate(output_path)

    assert output_path.exists()
    doc = Document(output_path)
    # Check for section headings
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert "1. Introduction" in headings
    assert "2. Protocol Description" in headings
    assert "3. List of Included Documents" in headings
    assert "4. Data Collection and Processing" in headings
    assert "5. SDTM Datasets" in headings
    assert "6. Data Conformance" in headings
    assert "7. Appendices" in headings
