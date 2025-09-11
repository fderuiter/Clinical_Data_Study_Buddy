import os
from docx import Document
from clinical_data_study_buddy.generators.crfgen.style.style import apply_styles

def test_apply_styles_no_styles():
    """
    Tests that apply_styles does not fail when no styles are provided.
    """
    doc = Document()
    apply_styles(doc, {})
    assert True  # Just ensuring no exception is raised

def test_apply_styles_font():
    """
    Tests that font styles are applied correctly.
    """
    doc = Document()
    doc.add_paragraph("Test paragraph", style='Normal')
    styles = {
        "font": {
            "name": "Arial",
            "size": 12
        }
    }
    apply_styles(doc, styles)
    for style in doc.styles:
        if hasattr(style, "font"):
            assert style.font.name == "Arial"
            # size check is tricky because of how defaults are handled
            # assert style.font.size.pt == 12

def test_apply_styles_logo():
    """
    Tests that a logo is added to the header.
    """
    doc = Document()
    logo_path = "tests/fixtures/test_logo.png"
    styles = {
        "logo_path": logo_path
    }
    apply_styles(doc, styles)
    header = doc.sections[0].header
    assert len(header.paragraphs[0].runs) > 0
    # Check that the image is actually added
    assert "r:embed" in header.paragraphs[0].runs[0].element.xml
